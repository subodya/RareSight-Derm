# -*- coding: utf-8 -*-
"""Generate CHAPTER-02.docx (Literature Review) in the FYP Final Report structure and the
author's voice. Restructures the prior draft to: 2.3 Problem Domain Review -> 2.4 Technological
Review -> 2.5 Existing Work -> 2.6 Evaluation & Benchmarking -> 2.7 Summary. Removes all
comparison with the project's own previously-attempted methods (per the writing-style directive).
Formatting per the marking guide: Times New Roman; chapter 16pt CAPS bold, H2 14pt bold,
H3 12pt bold, body 12pt 1.5 justified."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
TARGET = os.path.join(ROOT, "CHAPTER-02.docx")
FALLBACK = os.path.join(ROOT, "CHAPTER-02_DRAFT.docx")
FONT = "Times New Roman"


def _set_font(style, size, bold=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def build():
    doc = Document()
    # base styles
    _set_font(doc.styles["Normal"], 12)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.5
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        _set_font(doc.styles[name], size, bold=True)
        doc.styles[name].paragraph_format.space_before = Pt(12)
        doc.styles[name].paragraph_format.space_after = Pt(6)

    def chap(text):
        p = doc.add_paragraph(text.upper(), style="Heading 1")
        return p

    def h2(text):
        doc.add_paragraph(text, style="Heading 2")

    def h3(text):
        doc.add_paragraph(text, style="Heading 3")

    def body(text):
        p = doc.add_paragraph(text, style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def figure(cap, note):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[" + note + "]")
        r.italic = True
        r.font.size = Pt(11)
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = c.add_run(cap)
        rc.italic = True
        rc.font.size = Pt(11)

    # ============================ CONTENT ============================
    chap("Chapter 02: Literature Review")

    h2("2.1 Chapter overview")
    body("This chapter critically analyses the body of knowledge on which RareSight is built and "
         "situates the project within it. It opens by mapping the research area and reviewing the "
         "problem domain — the diagnosis of rare and underrepresented dermatological conditions "
         "under extreme data scarcity — before examining the technologies that compose a medical "
         "few-shot pipeline stage by stage. It then analyses the existing work through thematic "
         "synthesis rather than paper-by-paper summary, evaluates how that work is benchmarked, and "
         "concludes by consolidating the gaps that motivate the methodology of Chapter 3. The "
         "literature surveyed is predominantly recent (2020–2025), with seminal earlier works "
         "(2017–2019) retained where they remain foundational, and each factual claim is attributed "
         "to an authoritative source.")

    h2("2.2 Concept map")
    figure("Figure 2.1 – Concept map of the research domain.",
           "INSERT figures/concept_map.png — research streams (medical VLMs, prompt learning, "
           "few-shot meta-learning, multi-modal/metadata fusion, trust mechanisms, rare-disease AI) "
           "converging on the integration gaps RareSight addresses")
    body("Figure 2.1 illustrates the conceptual structure of the domain; the detailed version is "
         "provided in Appendix I. The map organises the literature into interconnected research "
         "streams that converge on rare-disease dermatology. The first stream traces medical "
         "vision-language models from general-domain CLIP to specialised BiomedCLIP and learnable "
         "prompts; the second charts few-shot meta-learning from Prototypical Networks through "
         "dermatological applications; the third contextualises rare-disease artificial intelligence "
         "(AI), with its clinical urgency and data scarcity; and a fourth, cross-cutting stream "
         "concerns the trust mechanisms — calibration, uncertainty, and abstention — that clinical "
         "deployment requires. At their convergence, three integration gaps emerge that frame the "
         "remainder of this chapter and directly motivate RareSight's training-free, multi-modal, "
         "and trust-aware design.")

    # ---------------- 2.3 PROBLEM DOMAIN REVIEW ----------------
    h2("2.3 Problem domain review")
    body("This section establishes the theoretical and contextual foundations of the research. It "
         "introduces the domain and its key terms, decomposes the problem into its principal "
         "challenges, surveys the frameworks and standards that shape how solutions are built and "
         "judged, and closes by positioning RareSight's conceptual architecture against those "
         "challenges as a bridge to the methodology.")

    h3("2.3.1 Introduction to the domain")
    body("The domain of this research is the automated diagnosis of skin disease from images under "
         "severe data scarcity. Rare diseases collectively affect an estimated 300 million people "
         "worldwide, around half of them children, yet each individual condition is by definition "
         "uncommon, so that most have fewer than fifty documented clinical images globally "
         "(Schaefer et al., 2020). Dermatology is an especially acute setting for this scarcity: "
         "common conditions such as melanoma are supported by tens of thousands of labelled images "
         "and attain accuracies above ninety percent, whereas rare genodermatoses remain data-poor "
         "and poorly served (Liopyris et al., 2022). The clinical consequence is the so-called "
         "diagnostic odyssey, a journey from first symptom to correct diagnosis that averages five "
         "to seven years and during which a large fraction of patients are misdiagnosed at least "
         "once (Schaefer et al., 2020).")
    body("Two technologies have emerged as the principal responses to this scarcity, and both are "
         "central to the domain. The first is few-shot meta-learning, in which a model learns to "
         "classify from only a handful of labelled examples per class — typically framed as an "
         "N-way K-shot episode, where the model must distinguish N classes given K support images "
         "of each — rather than from thousands of examples (Snell et al., 2017). The second is the "
         "vision-language model (VLM), which aligns images and text in a shared embedding space "
         "through contrastive pre-training and can therefore classify by comparing an image to "
         "natural-language class descriptions, even for categories unseen during training "
         "(Radford et al., 2021). A further distinction structures the dermatological domain "
         "specifically: dermoscopy images, captured through a contact lens under controlled "
         "illumination, differ markedly from clinical photographs taken with a smartphone under "
         "uncontrolled lighting (Tschandl et al., 2018; Pacheco et al., 2020). The interaction of "
         "scarcity, alignment, and this modality distinction defines the problem this thesis "
         "addresses.")
    body("Dermatology is a particularly suitable proving ground for these technologies for three "
         "reasons. It is image-centric, so a visual model can in principle capture the diagnostic "
         "signal a clinician uses; it records structured metadata as a matter of routine, offering "
         "complementary information at no additional cost; and it exhibits a pronounced long tail, "
         "in which a few common conditions dominate the data while thousands of rarer dermatoses "
         "are scarcely represented (Liopyris et al., 2022). The term underrepresented here is "
         "deliberately broad: it encompasses not only rare diagnoses but also populations and skin "
         "tones that are under-served by existing datasets, a fairness dimension that is acute in a "
         "low-resource deployment context where teledermatology may be the only realistic route to "
         "specialist opinion (Schaefer et al., 2020). The remainder of this section decomposes the "
         "challenges these characteristics create.")

    h3("2.3.2 Key challenges and sub-areas")
    body("The problem domain can be decomposed into four challenges that recur throughout the "
         "literature and that any deployable rare-disease diagnostic system must confront.")
    body("Extreme data scarcity and long-tailed distributions. The defining difficulty is that "
         "deep learning conventionally requires upwards of a thousand examples per class, while "
         "rare conditions supply only a handful (Schaefer et al., 2020). Conventional transfer "
         "learning — pre-training on a large common-disease archive and fine-tuning on rare classes "
         "— offers limited benefit when a rare disease exhibits a novel phenotype with no visually "
         "similar common counterpart (Liopyris et al., 2022). Scarcity therefore cannot be "
         "engineered away by more training and instead demands methods that generalise from minimal "
         "evidence, which motivates the few-shot framing of the next sub-area.")
    body("Domain shift between acquisition modalities. A model that performs well on curated "
         "dermoscopy must still contend with the substantial distribution shift introduced by "
         "real-world clinical photography, where focus, lighting, scale, and background vary widely "
         "(Pacheco et al., 2020). This shift is precisely the condition under which a triage tool "
         "would operate in a low-resource setting, where smartphone capture is the norm, and it is "
         "frequently the point at which laboratory accuracy fails to transfer (Pachetti & "
         "Colantonio, 2024). Robustness to modality shift is thus a first-class requirement rather "
         "than an afterthought, and it leads naturally to the question of what additional "
         "information can stabilise prediction when the image alone is unreliable.")
    body("Under-exploitation of complementary modalities. Most few-shot medical pipelines operate "
         "on images alone, discarding the clinical text and structured patient metadata that are "
         "routinely recorded and are strongly diagnostic (Mahajan et al., 2020; Özdemir et al., "
         "2025). Disease names and symptom descriptions carry visual semantics, and patient age, "
         "sex, and lesion site materially shift diagnostic priors, yet this side-information is "
         "rarely aggregated into the prototype itself (Pachetti & Colantonio, 2024). Because "
         "auxiliary signal should be most valuable exactly where visual examples are scarcest, its "
         "neglect is most costly in the rare-disease regime, which is the regime this research "
         "targets.")
    body("Clinical trust: calibration, uncertainty, and safe abstention. A system intended for "
         "decision support must communicate trustworthy confidence and must fail safely on inputs "
         "outside its competence. Modern networks are systematically over-confident, so raw "
         "softmax scores misrepresent reliability unless explicitly calibrated (Guo et al., 2017), "
         "and a triage tool deployed beyond its validated domain will inevitably meet "
         "out-of-distribution inputs it should decline rather than misclassify (Rajaraman et al., "
         "2022). These trust properties are essential for deployment yet, as the existing-work "
         "review will show, are seldom reported together, completing the set of challenges that "
         "frame RareSight's design.")

    h3("2.3.3 Existing frameworks, standards, and domain approaches")
    body("Research in this domain is shaped by several de facto standards that constrain how "
         "systems are designed and evaluated. The dominant evaluation framework is the N-way K-shot "
         "episodic protocol introduced for meta-learning, under which accuracy is averaged over "
         "many randomly sampled tasks and reported with confidence intervals across seeds (Snell et "
         "al., 2017; Finn et al., 2017). Dermatological research is further organised around a small "
         "number of canonical data resources that function as community benchmarks: the "
         "International Skin Imaging Collaboration (ISIC) archive and the HAM10000 dermoscopy "
         "collection (Tschandl et al., 2018), and, for clinical photography with metadata, "
         "PAD-UFES-20 (Pacheco et al., 2020). These resources standardise comparison but also "
         "constrain it, because their label taxonomies and acquisition conditions become implicit "
         "assumptions of any model trained on them.")
    body("For clinical trustworthiness, a parallel set of reporting standards has emerged. "
         "Calibration is quantified by the expected calibration error (ECE) and the Brier score, "
         "and threshold-independent discrimination by the area under the receiver operating "
         "characteristic curve (AUROC), with temperature scaling established as a strong, minimal "
         "calibration baseline (Guo et al., 2017; Rajaraman et al., 2022). The incumbent solution "
         "approach against which new methods are measured remains supervised convolutional "
         "classification, optionally with transfer learning and augmentation, which excels on "
         "common conditions but degrades under scarcity and novel phenotypes (Liopyris et al., "
         "2022). Finally, professional and regulatory frameworks — the British Computer Society "
         "Code of Conduct and the General Data Protection Regulation — govern data handling, "
         "fairness, and accountability for any clinically oriented system, and their implications "
         "for RareSight are examined in Chapter 4. Together these standards define both the playing "
         "field and the obligations within which the proposed architecture must operate.")

    h3("2.3.4 Proposed architecture (positioning)")
    figure("Figure 2.2 – Proposed conceptual architecture of RareSight.",
           "INSERT figures/architecture.png — frozen BiomedCLIP encoder feeding an aligned-prototype "
           "blend, optional CoOp prompt, metadata-likelihood fusion, calibration and open-set "
           "guard, with modality routing to a resolution-banded dermoscopy path and a clinical path")
    body("Figure 2.2 presents RareSight's high-level conceptual architecture, positioned directly "
         "against the four challenges above. At its core is a frozen BiomedCLIP encoder — a "
         "ViT-B/16 vision transformer aligned with a PubMedBERT text encoder — which supplies "
         "medical domain knowledge without any task-specific training and thereby answers the "
         "scarcity challenge by classifying from a few aligned prototypes rather than from a large "
         "training set (Zhang et al., 2025). Few-shot inference is performed in this frozen aligned "
         "space: class prototypes are formed by blending image embeddings with a text-prompt "
         "ensemble and a modality-gap correction, an approach that exploits the semantic content of "
         "class descriptions while leaving the backbone untouched, in keeping with prompt-tuning "
         "evidence that minimal adaptation can outperform heavier fine-tuning (Zhou et al., 2022). "
         "An optional, compact learned prompt provides a small amount of task adaptation at "
         "negligible parameter cost.")
    body("The remaining components address the complementary-modality and trust challenges "
         "respectively. Structured patient metadata is incorporated as a class-conditional "
         "likelihood combined with the visual posterior, a lightweight and training-free mechanism "
         "intended to supply signal where the image is uninformative. Temperature-scaled "
         "calibration and a Mahalanobis-distance open-set guard provide trustworthy confidence and "
         "safe abstention, while a modality router separates a resolution-banded dermoscopy path "
         "from a dedicated clinical-photography path so that smartphone inputs are matched to "
         "prototypes built under their own acquisition conditions. This conceptual design — frozen "
         "alignment, multi-modal prototypes, and integrated trust mechanisms — is formalised as a "
         "methodology in Chapter 3.")

    # ---------------- 2.4 TECHNOLOGICAL REVIEW ----------------
    h2("2.4 Technological review")
    body("This section examines the technologies that compose a medical few-shot pipeline stage by "
         "stage — data acquisition and pre-processing, model architectures and representations, "
         "adaptation and inference mechanisms, and explainability tooling — emphasising critical "
         "analysis of the trade-offs rather than enumeration.")

    h3("2.4.1 Data acquisition and pre-processing")
    body("Dermatological data span a spectrum from curated dermoscopy to uncontrolled clinical "
         "photography, and the contrast is methodologically consequential. HAM10000 provides "
         "standardised dermoscopic images with verified diagnostic labels (Tschandl et al., 2018), "
         "whereas PAD-UFES-20 supplies real smartphone photographs accompanied by structured "
         "patient metadata (Pacheco et al., 2020); evaluating across both therefore constitutes a "
         "genuine test of generalisation rather than within-distribution recall. Standard "
         "pre-processing resizes images to the encoder's expected resolution and normalises them to "
         "the pre-training statistics, and for an aligned VLM the preservation of native image "
         "fidelity matters, since full-resolution capture retains the fine textural cues — lesion "
         "borders, scale, and pigment network — that aggressive down-sampling discards. Two "
         "augmentation techniques are particularly relevant to clinical photography: colour "
         "constancy normalisation, which removes illuminant bias and has long improved skin-lesion "
         "classification (Barata et al., 2015), and test-time augmentation, in which conservative, "
         "label-preserving transforms such as flips and rotations expand a small support set into "
         "virtual examples without any retraining (Shorten & Khoshgoftaar, 2019). Critically, the "
         "value of such transforms is contingent on the encoder and modality, so they must be "
         "validated rather than assumed. A colour normalisation that benefits a convolutional model "
         "trained on raw photographs may, for instance, perturb the inputs of a contrastively "
         "pre-trained encoder away from the statistics it expects, illustrating that pre-processing "
         "interacts with the backbone and cannot be transferred uncritically between systems. "
         "Resolution is a further pre-processing axis with diagnostic consequences, because fine "
         "texture is concentrated at high spatial frequencies and the band of resolutions at which "
         "a lesion is presented to the encoder can determine whether the cue survives encoding, "
         "motivating resolution-aware handling rather than a single fixed input size (Tschandl et "
         "al., 2018).")

    h3("2.4.2 Model architectures and feature representations")
    body("Medical image analysis has transitioned from handcrafted descriptors to learned "
         "representations, and within learned representations from convolutional to attention-based "
         "backbones. Convolutional networks such as ResNet remain common in few-shot dermatology "
         "(Mahajan et al., 2020), but Vision Transformers encode images as sequences of patches and "
         "support the global reasoning that contrastive alignment exploits (Dosovitskiy et al., "
         "2021). The decisive architectural development for this domain is the contrastive "
         "vision-language model: CLIP demonstrated that training on large image–text corpora yields "
         "a shared embedding space enabling zero-shot transfer (Radford et al., 2021), and "
         "BiomedCLIP specialised this recipe to fifteen million biomedical image–text pairs across "
         "dermatology, pathology, radiology, and ophthalmology, pairing a ViT-B/16 vision encoder "
         "with a PubMedBERT text encoder (Zhang et al., 2025). A more recent line of work pursues "
         "dermatology-specific foundation models pre-trained on large multi-source skin-image "
         "collections spanning several imaging modalities (Yan et al., 2025). The critical property "
         "uniting these models is alignment: their diagnostic power derives from a carefully "
         "structured image–text space, which makes the preservation of that alignment during "
         "adaptation a central technical concern.")
    body("The choice between convolutional and transformer backbones is itself a trade-off under "
         "scarcity. Vision Transformers are data-hungry to train from scratch, which would ordinarily "
         "disadvantage them in low-data settings, but when a transformer is used as a frozen, "
         "large-scale pre-trained encoder this weakness is neutralised and its capacity for global "
         "reasoning is retained (Dosovitskiy et al., 2021). A subtler property of contrastive "
         "image–text spaces is the modality gap: image and text embeddings occupy separate cones of "
         "the shared space rather than coinciding, so a naive comparison of an image to a text "
         "prototype is biased unless the gap is accounted for (Liang et al., 2022). This phenomenon "
         "is directly relevant to any method that blends visual and textual prototypes, and it "
         "informs the modality-gap correction in the proposed architecture.")

    h3("2.4.3 Adaptation and few-shot inference")
    body("Given a frozen aligned encoder, several families of technique adapt it to a downstream "
         "few-shot task without full fine-tuning. Prompt learning replaces hand-crafted text "
         "prompts with a small set of continuous, learnable context vectors optimised by gradient "
         "descent, improving CLIP few-shot accuracy while leaving the backbone frozen and adapting "
         "only a few thousand parameters (Zhou et al., 2022). A complementary, training-free "
         "strategy enriches prompts rather than learning them, aggregating many language-model-"
         "generated descriptions of each class into a more robust textual prototype. Adapter "
         "methods occupy a middle ground: training-free cache models such as Tip-Adapter store "
         "support features as keys and labels as values and classify by feature retrieval, blending "
         "a non-parametric cache with the zero-shot prediction (Zhang et al., 2022). Prototype-based "
         "inference itself involves consequential choices — the distance metric (cosine versus "
         "Euclidean), the prototype-construction rule (mean versus weighted pooling), and the "
         "fusion strategy for combining modalities, which ranges from simple concatenation through "
         "gated fusion to cross-modal attention (Shakeri et al., 2024). The recurring tension across "
         "these techniques is between expressive capacity and the integrity of the aligned space: "
         "approaches that add trainable capacity risk overfitting under K-shot scarcity, whereas "
         "minimal, alignment-preserving adaptation tends to be more robust (Zhou et al., 2022). "
         "This tension is sharpened by the small support sets characteristic of rare-disease "
         "diagnosis: with only a few images per class, any component that learns visual–textual or "
         "visual–metadata correlations has little data from which to estimate them and may instead "
         "fit spurious patterns, so the safer design injects auxiliary signal through fixed, "
         "estimated, or retrieval-based mechanisms rather than through additional learned layers "
         "(Pachetti & Colantonio, 2024). The hyperparameters that govern these mechanisms — the "
         "number of shots, the metric temperature, and the relative weight of image, text, and "
         "metadata evidence — therefore exert a strong influence on performance and warrant explicit "
         "study rather than default settings.")

    h3("2.4.4 Explainability and evaluation tooling")
    body("Interpretability tooling must match the backbone. For convolutional networks, "
         "gradient-weighted class activation mapping (Grad-CAM) is standard, but for transformer "
         "encoders the appropriate attribution method is Attention Rollout, which composes the "
         "attention matrices across layers to trace how the classification token aggregates "
         "evidence from image patches (Abnar & Zuidema, 2020; Suara et al., 2024). Selecting the "
         "transformer-appropriate method is not a cosmetic choice, since applying a convolutional "
         "attribution to a patch-based encoder yields misleading saliency, and the faithfulness of "
         "an explanation — whether the highlighted region genuinely drives the prediction — matters "
         "more in a clinical setting than visual plausibility alone. On the measurement side, the field "
         "employs a consistent toolkit of metrics — accuracy with confidence intervals, "
         "macro-averaged F1 for class imbalance, AUROC for discrimination, and ECE and the Brier "
         "score for calibration — which the evaluation section examines in detail. The maturity of "
         "this tooling means that a credible system is judged not only on accuracy but on "
         "interpretable evidence and calibrated confidence, reinforcing the trust requirements "
         "identified in the problem domain.")

    # ---------------- 2.5 EXISTING WORK ----------------
    h2("2.5 Existing work")
    body("This section analyses the existing literature through thematic synthesis, discussing "
         "multiple works together under each theme to surface trends and divergences, before "
         "linking the themes to consolidate the gaps that motivate this research.")

    h3("2.5.1 Thematic grouping")
    body("Medical vision-language models. The evolution of medical VLMs marks a shift from "
         "supervised learning toward self-supervised pre-training on image–text pairs. After CLIP "
         "established contrastive transfer (Radford et al., 2021), domain-specific adaptations "
         "followed: BioViL for radiology, trained on chest X-ray and report pairs (Boecking et al., "
         "2022); PMC-CLIP, built from PubMed Central figure–caption pairs (Lin et al., 2023); and "
         "BiomedCLIP, trained across multiple specialties to state-of-the-art zero-shot performance "
         "(Zhang et al., 2025). However, this stream evaluates almost exclusively in zero-shot or "
         "fully-supervised settings; a survey of seventy-three CLIP medical-imaging papers found "
         "that, while many adapted CLIP for medical tasks, none combined a medical VLM with N-way "
         "K-shot episodic adaptation (Zhao et al., 2025).")
    body("Prompt learning for vision-language models. Because a VLM conditions classification on "
         "text, prompt design materially affects accuracy. CoOp introduced learnable continuous "
         "prompts that improve few-shot accuracy at negligible parameter cost (Zhou et al., 2022), "
         "while training-free prompt-ensemble methods avoid optimisation altogether by aggregating "
         "many descriptive phrasings per class. A closely related family augments the frozen encoder "
         "with lightweight adapters: training-free cache models such as Tip-Adapter classify by "
         "retrieving similar support features and require no gradient updates, embodying the same "
         "principle that substantial gains are available without disturbing the backbone (Zhang et "
         "al., 2022). However, this literature is evaluated almost entirely on natural-image "
         "benchmarks, leaving its behaviour on aligned biomedical models, its sensitivity to class "
         "imbalance, and its interaction with structured patient metadata underexplored.")
    body("Few-shot meta-learning for medical imaging. Prototypical Networks classify by distance "
         "to class-mean embeddings and remain attractive for their simplicity and interpretability "
         "relative to gradient-based meta-learners such as MAML, which are costlier and less stable "
         "(Snell et al., 2017; Finn et al., 2017; Singh et al., 2021). Dermatological adaptations "
         "include Meta-DermDiagnosis, which applied Prototypical Networks to skin-disease "
         "classification (Mahajan et al., 2020), difficulty-aware scheduling for rare diseases (Li "
         "et al., 2020), and an episodic ResNet-50 pipeline with ResizeMix augmentation reaching "
         "seventy-eight percent five-shot accuracy on the SD-198 dataset, whose authors explicitly "
         "name the integration of vision-language models with episodic meta-learning as essential "
         "future work (Özdemir et al., 2025). A systematic review of one hundred and twenty-seven "
         "few-shot medical-imaging papers confirms the pattern: only a minority use episodic "
         "training, and none integrate vision-language pre-training (Pachetti & Colantonio, 2024).")
    body("Multi-modal and metadata fusion. As medical AI moves beyond image-only analysis, "
         "multi-modal masked autoencoders have jointly modelled chest X-rays and reports (Chen et "
         "al., 2022) and Med-Flamingo has adapted in-context learning to medical visual question "
         "answering (Moor et al., 2023). Learnable multi-modal prototypes have also been explored "
         "for few-shot classification (Shakeri et al., 2024). However, this body of work targets "
         "learned fusion in large-scale supervised settings and rarely incorporates the structured "
         "metadata — age, sex, and anatomical site — that dermatology records as standard. The "
         "distinction matters because in-context and learned-fusion approaches presuppose either "
         "many examples or a trainable fusion module, neither of which is well suited to a "
         "prototype-based pipeline operating on a few support images, where a probabilistic "
         "treatment of metadata as a class-conditional prior is both lighter and more robust "
         "(Moor et al., 2023).")
    body("Calibration, uncertainty, and out-of-distribution detection. Modern networks are "
         "systematically over-confident, and temperature scaling substantially improves calibration "
         "as measured by ECE (Guo et al., 2017); subsequent work has extended calibration and "
         "uncertainty analysis to class-imbalanced clinical imaging (Rajaraman et al., 2022; "
         "Lambert et al., 2022). Distance-based out-of-distribution detection, including Mahalanobis "
         "scoring in feature space, enables a model to abstain rather than misclassify. However, "
         "these trust mechanisms are rarely reported alongside few-shot accuracy.")
    body("Rare-disease AI systems. Research targeting rare diseases specifically remains nascent: a "
         "scoping review of one hundred and fifty-nine machine-learning rare-disease papers found "
         "that only a small fraction used few-shot learning and none used vision-language models "
         "(Schaefer et al., 2020). Most such systems rely on supervised learning with augmentation "
         "or transfer from related common diseases, which falters when a rare disease has no "
         "visually similar common counterpart (Liopyris et al., 2022).")

    h3("2.5.2 Cross-theme linking and identified gaps")
    body("Read together, these themes reveal a field that has advanced rapidly but divergently. "
         "Vision-language researchers evaluate on zero-shot or supervised benchmarks, meta-learning "
         "researchers rely on convolutional backbones without medical pre-training, and trust "
         "mechanisms are treated as an afterthought rather than a design requirement — a "
         "fragmentation in which the complementary strengths of each stream are seldom combined "
         "(Zhao et al., 2025; Pachetti & Colantonio, 2024). The overlap between streams is "
         "instructive: works that touch two themes, such as learnable multi-modal prototypes "
         "(Shakeri et al., 2024) or in-context medical VLMs (Moor et al., 2023), still stop short of "
         "prototype-based few-shot inference that aggregates auxiliary metadata under genuine "
         "scarcity. From this synthesis three gaps consolidate. First, medical VLMs are seldom "
         "integrated with few-shot adaptation, and the fragility of their contrastive alignment "
         "under heavier adaptation is neither tested nor reported. Second, few-shot prototypes are "
         "predominantly image-only, ignoring text-prompt ensembles and structured patient metadata. "
         "Third, evaluation relies on simulated rarity and omits the calibration, uncertainty, and "
         "abstention that deployment demands. These three gaps define the design targets that the "
         "remainder of the thesis pursues.")

    # ---------------- 2.6 EVALUATION AND BENCHMARKING ----------------
    h2("2.6 Evaluation and benchmarking")
    body("This section critically assesses how work in the domain is evaluated — the datasets used, "
         "the metrics and protocols applied, and the limitations that recur — in order to justify "
         "the evaluation strategy adopted in later chapters.")

    h3("2.6.1 Datasets and their scope")
    body("Dermatological few-shot research draws on a small set of resources of differing scope and "
         "diversity. HAM10000 offers around ten thousand curated dermoscopic images across seven "
         "diagnostic categories and is well suited to statistically powered evaluation (Tschandl et "
         "al., 2018); PAD-UFES-20 contributes roughly two thousand clinical smartphone photographs "
         "with per-lesion metadata, providing the domain shift and auxiliary information absent from "
         "dermoscopy benchmarks (Pacheco et al., 2020); and SD-198 supplies a broad clinical "
         "taxonomy for fine-grained recognition (Özdemir et al., 2025). A persistent weakness is the "
         "reliance on simulated rarity, in which minority classes are sub-sampled from a large "
         "dataset and therefore retain its acquisition conditions, overstating generalisation (Li "
         "et al., 2020; Pachetti & Colantonio, 2024). A further limitation is skin-tone diversity: "
         "widely used benchmarks under-represent darker skin types, which raises fairness concerns "
         "for any deployed triage system and is revisited in Chapter 4. Availability also varies: "
         "dermoscopy archives are large and openly licensed, whereas clinical-photograph datasets "
         "with consented metadata are smaller and harder to obtain, which partly explains why the "
         "literature gravitates toward dermoscopy despite clinical photography being the modality a "
         "low-resource triage tool would actually encounter (Pacheco et al., 2020). A defensible "
         "evaluation therefore pairs a large dermoscopy benchmark for statistically powered "
         "measurement with an independent clinical dataset for genuine domain transfer, rather than "
         "relying on either alone.")

    h3("2.6.2 Metrics and protocols")
    body("Evaluation in the domain combines standard and trust-specific metrics. Accuracy remains "
         "primary and is reported as a mean with confidence intervals over many episodes and "
         "several seeds, following the N-way K-shot protocol (Snell et al., 2017). Macro-averaged "
         "F1 balances precision and recall under class imbalance, AUROC measures threshold-"
         "independent discrimination, and top-K accuracy reflects the differential-diagnosis nature "
         "of clinical decisions. For trust, ECE and the Brier score quantify confidence reliability "
         "(Guo et al., 2017). Beyond single-dataset accuracy, two protocols are especially "
         "informative for scarcity: shot-scaling analysis, which characterises performance as the "
         "number of support examples varies, and cross-dataset evaluation, which measures transfer "
         "under domain shift. Appropriate baselines for an aligned-VLM few-shot system include "
         "zero-shot BiomedCLIP, a convolutional Prototypical Network that isolates the value of "
         "medical pre-training, and a fully-supervised reference; the last must be read as a "
         "reference point rather than a competitor, because a figure computed under a different "
         "protocol — for example seven-way supervised classification — is not directly comparable "
         "with five-way few-shot accuracy.")

    h3("2.6.3 Limitations in existing evaluation")
    body("Several limitations recur across the literature and shape the evaluation design of this "
         "thesis. Inconsistent test splits hinder direct comparison between studies; many papers "
         "omit component-wise ablations, leaving unclear which design choices drive performance; "
         "few report uncertainty quantification or calibrated confidence; error analysis is often "
         "shallow; and simulated rarity inflates apparent generalisation (Pachetti & Colantonio, "
         "2024). The appropriate response, adopted in later chapters, is to report results with "
         "paired significance tests, to ablate components individually, to include calibration and "
         "uncertainty metrics, and to evaluate across an independent clinical dataset rather than on "
         "artificially restricted classes — while keeping comparisons within matched protocols so "
         "that reported gains are meaningful. The last point deserves emphasis: because few-shot "
         "and fully-supervised results are produced under different numbers of classes and "
         "different amounts of training data, placing them in a single comparison table without "
         "qualification can misrepresent both, and a rigorous evaluation states the protocol "
         "explicitly and compares only like with like (Pachetti & Colantonio, 2024). Adhering to "
         "these principles is what allows a literature otherwise fragmented across protocols to be "
         "synthesised responsibly, and it is the standard the evaluation chapters of this thesis "
         "adopt.")

    # ---------------- 2.7 SUMMARY ----------------
    h2("2.7 Chapter summary")
    body("This chapter reviewed the literature underpinning RareSight across medical vision-language "
         "models, prompt learning, few-shot meta-learning, multi-modal and metadata fusion, "
         "calibration and out-of-distribution detection, and rare-disease AI, and assessed how that "
         "work is benchmarked. Four insights emerged. First, medical VLMs achieve strong zero-shot "
         "performance but are rarely adapted for few-shot diagnosis, and the literature does not "
         "test the robustness of their aligned space to adaptation. Second, prompt-learning evidence "
         "indicates that minimal adaptation can outperform heavier fine-tuning, motivating a "
         "training-free or near-training-free design. Third, complementary modalities — especially "
         "structured patient metadata — are systematically under-exploited in few-shot work. Fourth, "
         "trust mechanisms such as calibration, uncertainty, and abstention are essential for "
         "deployment yet seldom reported together, and evaluation too often relies on simulated "
         "rarity. RareSight is positioned to address these gaps by integrating BiomedCLIP's medical "
         "knowledge with training-free multi-modal prototypes, class-conditional metadata fusion, "
         "calibration, and open-set rejection, and by evaluating on full-resolution HAM10000 and the "
         "independent PAD-UFES-20 dataset. The next chapter formalises this positioning as a "
         "methodology.")

    # ---- save (handle the file being open in Word) ----
    out = TARGET
    try:
        doc.save(TARGET)
    except PermissionError:
        out = FALLBACK
        doc.save(FALLBACK)
    print("SAVED:", out)
    print("paragraphs:", len(doc.paragraphs))


if __name__ == "__main__":
    build()
